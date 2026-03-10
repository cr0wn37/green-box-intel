import os
import uuid
import time
import tempfile
import io
from pathlib import Path
from dotenv import load_dotenv
import json
import re
from pypdf import PdfReader
import fitz
import gc
import boto3
from supabase import create_client, Client
from database import supabase

from fastapi import FastAPI, UploadFile, File, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import Response
from fastapi.middleware.cors import CORSMiddleware
from typing import List
from datetime import datetime
from fastapi import Form
from pydantic import BaseModel
from database import DatabaseManager

# AI & Security Imports
from groq import Groq
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from fastapi.responses import StreamingResponse
from presidio_anonymizer.entities import OperatorConfig

# Your custom AWS engine
from aws_utils import (
    process_large_legal_pdf, 
    upload_safe_text_to_s3,
    s3_client,  
    BUCKET_NAME 
)

load_dotenv()

class RenameRequest(BaseModel):
    new_name: str
    user_id: str

# Initialize Clients
app = FastAPI(title="Green Box Legal - Intelligence Infrastructure")
client = Groq(api_key=os.getenv("GROQ_API_KEY"))
analyzer = None
anonymizer = None

app = FastAPI()

# Allow your Streamlit frontend to talk to this API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with your specific URL
    allow_methods=["*"],
    allow_credentials=True,
    allow_headers=["*"],
)

def get_pii_engines():
    """Wakes up the heavy spaCy models ONLY when a file is processed."""
    global analyzer, anonymizer
    
    if analyzer is not None and anonymizer is not None:
        return analyzer, anonymizer
        
    print("⏳ Booting up Presidio & spaCy into RAM...")
    from presidio_analyzer import AnalyzerEngine
    from presidio_anonymizer import AnonymizerEngine
    from presidio_analyzer.nlp_engine import SpacyNlpEngine
    
    # Force the 12MB small model instead of the 800MB large model
    configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
    }
    nlp_engine = SpacyNlpEngine(models=configuration["models"])
    
    analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
    anonymizer = AnonymizerEngine()
    
    return analyzer, anonymizer

# Global memory to track long-running jobs
# In production, swap this for Redis or a Database
processing_jobs = {}

# --- HELPER: WORD GENERATION ---
# --- HELPER: WORD GENERATION ---
def create_word_doc_stream(text_content, job_id, filename):
    doc = Document()
    
    # 1. NEW: Professional Meta-Header
    header = doc.sections[0].header
    htable = header.add_table(1, 2, doc.sections[0].page_width)
    htab_cells = htable.rows[0].cells
    htab_cells[0].text = "GREEN BOX LEGAL\nMedical Intelligence Report"
    htab_cells[1].paragraphs[0].text = f"Case ID: {job_id[:8]}\nSource: {filename}"
    htab_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 2. Existing "Privileged" title but centered
    title = doc.add_heading('PRIVILEGED & CONFIDENTIAL', 3)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Medical Chronology Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. Smart Parsing with Section Support
    lines = text_content.split('\n')
    table_data = []
    
    for line in lines:
        stripped_line = line.strip().replace('**', '')
        
        # Detect Sections for Styling
        if 'EXECUTIVE CASE SUMMARY' in stripped_line.upper():
            doc.add_heading('Executive Overview', level=2)
            continue
        elif 'MEDICAL CHRONOLOGY TABLE' in stripped_line.upper():
            doc.add_heading('Detailed Treatment Timeline', level=2)
            continue
        elif 'INJURY HIGHLIGHTS' in stripped_line.upper():
            if table_data:
                create_native_table(doc, table_data)
                table_data = []
            doc.add_heading('Key Injury Highlights', level=2)
            continue

        # Your original Smart Table detection
        if '|' in line:
            cells = [c.strip().replace('**', '') for c in line.split('|') if c.strip()]
            if cells and not all(set(c) <= {'-', ':', ' '} for c in cells):
                table_data.append(cells)
        else:
            if table_data:
                create_native_table(doc, table_data)
                table_data = []
            
            if stripped_line:
                if stripped_line.startswith('*') or stripped_line.startswith('-'):
                    doc.add_paragraph(stripped_line.lstrip('*- ').strip(), style='List Bullet')
                else:
                    doc.add_paragraph(stripped_line)

    if table_data:
        create_native_table(doc, table_data)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream



# Now you can use it like this:
def get_user_cases(user_id):
    response = supabase.table('cases').select("*").eq('user_id', user_id).execute()
    return response.data

db_manager = DatabaseManager()

def save_to_db(job_id: str, data: dict, user_id: str):
    """
    Replaces JSON saving with Supabase upsert.
    Bundles extra fields into the metadata dictionary.
    """
    try:
        # 1. Extract core fields from the 'data' dictionary provided
        case_name = data.get('case_name') or data.get('filename', 'Unnamed Case')
        chronology = data.get('chronology', '')
        pages = data.get('pages', 0)

        # 2. Bundle extra fields into the 'metadata' bucket
        # We use data.get() to safely pull values from the input dictionary
        metadata = {
            "filename": data.get("filename"),
            "timestamp": data.get("timestamp"),
            "status": data.get("status", "Completed"),
            "total_damages": data.get('total_damages', 0.0), 
            "file_list": data.get('file_list', []),
            "last_edited_by": user_id
        }

        # 3. Call the manager with the correct arguments
        db_manager.save_case(
            user_id=user_id,
            job_id=job_id,
            case_name=case_name,
            chronology=chronology,
            total_pages=pages,  # Matches the new argument name in DatabaseManager
            metadata=metadata
        )
        print(f"✅ Case {job_id} synced to Supabase (extra fields stored in metadata).")
        
    except Exception as e:
        print(f"❌ Supabase Save Error: {e}")

USAGE_DB = "user_usage.json"

def update_lifetime_usage(user_id: str, pages_to_add: int):
    """
    Subtracts processed pages from the user's 1500-page quota in Supabase.
    """
    try:
        db_manager.update_user_quota(user_id, pages_to_add)
        
        # Fetch the new remaining amount to return it
        profile = db_manager.get_user_profile(user_id)
        return profile['remaining_quota']
    except Exception as e:
        print(f"❌ Quota Update Error: {e}")
        return 0

def get_lifetime_usage(user_id: str):
    """
    Returns the current remaining quota for the user.
    """
    try:
        profile = db_manager.get_user_profile(user_id)
        return profile.get('remaining_quota', 0)
    except Exception as e:
        print(f"❌ Fetch Quota Error: {e}")
        return 0

def calculate_total_billing(chronology_text):
    # Added \s* in case Claude adds extra spaces before the bracket
    matches = re.findall(r"TOTAL_LIST:\s*\[(.*?)\]", chronology_text) 
    
    if matches:
        numbers_str = matches[-1]
        total = 0.0
        
        # Split the string by commas
        raw_items = numbers_str.split(',')
        
        for item in raw_items:
            # 1. Strip whitespace and remove accidental dollar signs
            clean_item = item.replace('$', '').replace(' ', '').strip()
            
            if clean_item:
                try:
                    # 2. Add to total item-by-item
                    total += float(clean_item)
                except ValueError:
                    # 3. If one item is text ("None", "TBD"), ignore it and keep summing the rest
                    pass 
                    
        return total
        
    return 0.0

def count_pdf_pages(file_content: bytes) -> int:
    """Accurately counts PDF pages using PyMuPDF."""
    try:
        # Open the PDF from memory
        doc = fitz.open(stream=file_content, filetype="pdf")
        page_count = len(doc)
        doc.close()
        return page_count
    except Exception as e:
        print(f"❌ Error counting pages: {e}")
        return 0

def create_native_table(doc, data):
    """Helper to build a real Word table from a list of lists with clean text"""
    if not data: return
    
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.autofit = True
    
    for r_idx, row_content in enumerate(data):
        for c_idx, cell_value in enumerate(row_content):
            if c_idx < len(table.columns):
                # CLEANING: Strip out the double stars here
                clean_value = cell_value.replace('**', '').strip()
                
                cell = table.cell(r_idx, c_idx)
                cell.text = clean_value
                
                # Make the header row bold using Word's native bolding
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def get_current_safe_text(job_id: str):
    """
    Retrieves the existing safe text. 
    Updated to be 'Cloud First' - skips local checks to ensure consistency.
    """
    try:
        # Direct fetch from S3 (Source of Truth)
        # This works for both AWS S3 and Supabase Storage (if configured with s3_client)
        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=f"safe-text/{job_id}_safe.txt")
        return response['Body'].read().decode('utf-8')
    except Exception as e:
        print(f"⚠️ Could not find safe text for {job_id}: {e}")
        return ""

def process_new_pdf(file_bytes, filename, job_id):
    """
    Standardized extraction and anonymization for single supplemental files.
    UPDATED: Features Page-by-Page Memory Chunking and Lazy Loading.
    """
    # 1. Create a safe temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        temp_path = tmp.name

    try:
        # 2. Extract Blocks via AWS Textract
        blocks = process_large_legal_pdf(temp_path)
        
        if not blocks:
            print(f"⚠️ No text found in {filename}")
            return ""

        # --- 3. CHUNKING PREP: Group lines by Page Number ---
        page_chunks = []
        current_page = None
        current_page_text = []

        for b in blocks:
            if b.get('BlockType') == 'LINE':
                page_num = b.get('Page', 1)
                text = b.get('Text', '')
                
                # If we hit a new page, save the old page and start a new one
                if current_page != page_num:
                    if current_page_text:
                        page_chunks.append("\n".join(current_page_text))
                    current_page = page_num
                    current_page_text = [f"--- SOURCE FILE: {filename} | PAGE: {page_num} ---"]
                
                current_page_text.append(text)
        
        # Don't forget to append the very last page
        if current_page_text:
            page_chunks.append("\n".join(current_page_text))

        # --- 4. APPLY UNIQUE TOKEN SHIELD (Page-by-Page) ---
        # Wake up the lazy-loaded PII engines
        local_analyzer, local_anonymizer = get_pii_engines()
        
        operators = {
            "PERSON": OperatorConfig("replace", {"new_value": "<PERSON>"}),
            "LOCATION": OperatorConfig("replace", {"new_value": "<LOCATION>"}),
            "PHONE_NUMBER": OperatorConfig("mask", {"masking_char": "*", "chars_to_mask": 10, "from_end": True})
        }
        
        safe_text_pieces = []
        
        for chunk in page_chunks:
            if len(chunk.strip()) > 0:
                results = local_analyzer.analyze(text=chunk, entities=["PERSON", "LOCATION", "PHONE_NUMBER"], language='en')
                anonymized = local_anonymizer.anonymize(text=chunk, analyzer_results=results, operators=operators)
                
                safe_text_pieces.append(anonymized.text)
                
                # GARBAGE COLLECTION: Keep RAM flat
                del results
                del anonymized
                gc.collect()
        
        # 5. Glue the safe pages back together
        return "\n".join(safe_text_pieces)

    except Exception as e:
        print(f"❌ Error processing supplemental PDF {filename}: {e}")
        return ""
        
    finally:
        # 6. Cleanup: Always remove the temp file
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except: 
                pass

bedrock_runtime = boto3.client(
    service_name='bedrock-runtime', 
    region_name=os.getenv('AWS_REGION', 'us-east-1')
)

# --- BACKGROUND WORKER ---
def run_intelligence_pipeline(job_id: str, temp_paths: list, file_names: str, total_pages: int, user_id: str):
    """
    Main AI Logic. 
    UPDATED: Features Page-by-Page Memory Chunking and Lazy Loading.
    """
    try:
        clean_file_names = [f.strip() for f in file_names.split(',')]
        page_chunks = [] # We will store text here grouped by page
        
        # 1. Loop through every uploaded file
        for path, original_name in zip(temp_paths, clean_file_names):
            if job_id in processing_jobs:
                processing_jobs[job_id]["status"] = f"Extracting: {original_name}..."
            
            print(f"DEBUG: Starting OCR for {job_id}") # Add this
            blocks = process_large_legal_pdf(path)
            print(f"DEBUG: OCR Complete. Found {len(blocks)} blocks") # Add this
            if not blocks:
                continue 
            
            # --- CHUNKING PREP: Group lines by Page Number ---
            current_page = None
            current_page_text = []
            
            for b in blocks:
                if b.get('BlockType') == 'LINE':
                    page_num = b.get('Page', 1)
                    text = b.get('Text', '')
                    
                    # If we hit a new page, save the old page and start a new one
                    if current_page != page_num:
                        if current_page_text:
                            page_chunks.append("\n".join(current_page_text))
                        current_page = page_num
                        current_page_text = [f"--- SOURCE FILE: {original_name} | PAGE: {page_num} ---"]
                    
                    current_page_text.append(text)
            
            # Don't forget to append the very last page
            if current_page_text:
                page_chunks.append("\n".join(current_page_text))
        
        if not page_chunks:
            if job_id in processing_jobs:
                processing_jobs[job_id]["status"] = "Failed: No text extracted"
            return

        # 2. Wake up the PII Engines (Lazy Load)
        if job_id in processing_jobs:
            processing_jobs[job_id]["status"] = "Applying Unique Token Shield..."
            
        local_analyzer, local_anonymizer = get_pii_engines()
        
        operators = {
            "PERSON": OperatorConfig("replace", {"new_value": "<PERSON_{index}>"}),
            "LOCATION": OperatorConfig("replace", {"new_value": "<LOCATION_{index}>"}),
            "PHONE_NUMBER": OperatorConfig("mask", {"masking_char": "*", "chars_to_mask": 10, "from_end": True})
        }
        
        safe_text_pieces = []
        total_redactions = 0
        
        # 3. THE LIFESAVER LOOP: Process one page at a time
        for chunk in page_chunks:
            if len(chunk.strip()) > 0:
                results = local_analyzer.analyze(text=chunk, entities=["PERSON", "LOCATION", "PHONE_NUMBER"], language='en')
                anonymized = local_anonymizer.anonymize(text=chunk, analyzer_results=results, operators=operators)
                
                safe_text_pieces.append(anonymized.text)
                total_redactions += len(results)
                
                # Manually dump the RAM after every single page
                del results
                del anonymized
                gc.collect()


        print("DEBUG: PII Redaction Complete. Sending to Claude...") # Add this
        # 4. Glue the safe pages back together
        safe_text = "\n".join(safe_text_pieces)

        # Upload safe text to S3
        upload_safe_text_to_s3(job_id, safe_text)

        # 5. Generate Chronology (Claude 4.6 Opus via AWS Bedrock)
        if job_id in processing_jobs:
            processing_jobs[job_id]["status"] = "Generating Unified Chronology (Claude 4.6)..."
            
        system_instructions = """You are a Senior Personal Injury Paralegal. 
                Format your response in Markdown with the following specific sections:
                # EXECUTIVE CASE SUMMARY
                [Provide a 3-5 sentence professional overview of the patient's injuries, major treatments, and the 'bottom line' of the case.]
                # MEDICAL CHRONOLOGY TABLE
                [1. Create a SINGLE unified Medical Chronology table from all provided text. 
                2. CRITICAL RULE: For the 'Source' column, you MUST use the exact filename and page number.
                3. Maintain consistency with tokens like <PERSON_1>.]
                # MEDICAL BILLING LEDGER 
                [ Create a MEDICAL BILLING LEDGER section. Extract every dollar amount associated with a date and provider. Columns: Date of Service, Provider, Description of Service, and Amount Billed.]
                # INJURY HIGHLIGHTS
                [List 3-5 key clinical findings with page references. (Only once at the end)]
                # TOTAL COST 
                [At the end of the report, provide a DATA BLOCK for the developer. List every numerical amount found in the billing ledger in this format: TOTAL_LIST: [120.00, 450.50, 1000.00]. Do not include currency symbols in the list.]
                """

        # Formulate the payload for Claude
        body = json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 8192, # Maximum output length
            "system": system_instructions, # Claude puts the system prompt here!
            "messages": [
                {"role": "user", "content": f"Redacted Content from Multiple Files:\n\n{safe_text}"}
            ],
            "temperature": 0
        })

        # Call the Bedrock API
        bedrock_response = bedrock_runtime.invoke_model(
            modelId="global.anthropic.claude-opus-4-6-v1", # The exact ID you provisioned
            body=body
        )
        
        # Parse the JSON response
        response_data = json.loads(bedrock_response.get('body').read())
        report_content = response_data.get('content')[0].get('text')

        total_val = calculate_total_billing(report_content)

        # 6. Prepare Final Data
        existing_job = processing_jobs.get(job_id, {})
        existing_case_name = existing_job.get("case_name", "Unnamed Case")
        original_filename = existing_job.get("filename", "Document")
        
        for path in temp_paths:
            temp_name = os.path.basename(path)
            report_content = report_content.replace(temp_name, original_filename)
        
        final_job_data = {
            "status": "Completed",
            "chronology": report_content,
            "total_damages": total_val,  
            "redactions_count": total_redactions, # Accurate count across all pages
            "pages": total_pages, 
            "file_list": [original_filename], 
            "timestamp": datetime.now().isoformat(),
            "has_chat": True,
            "case_name": existing_case_name
        }

        # Update Memory
        if job_id in processing_jobs:
            processing_jobs[job_id].update(final_job_data)

        # 7. SAVE TO SUPABASE
        save_to_db(job_id, final_job_data, user_id)
        
    except Exception as e:
        print(f"PIPELINE ERROR: {e}")
        if job_id in processing_jobs:
            processing_jobs[job_id]["status"] = f"Failed: {str(e)}"
    finally:
        # Cleanup temp files
        for path in temp_paths:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except: pass

@app.post("/process-intelligence")
async def start_processing(
    background_tasks: BackgroundTasks, 
    files: List[UploadFile] = File(...), 
    case_name: str = Form(None),
    user_id: str = Form(...)  # <--- NEW: Required to link data to the user
):
    print(f"👉 DEBUG: New Job for User: {user_id} | Case: {case_name}")

    # 1. Validate PDF format
    for file in files:
        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail=f"File {file.filename} is not a PDF.")
    
    job_id = str(uuid.uuid4())
    temp_paths = []
    total_pages = 0 

    try:
        # 2. Save files safely (Cloud Compatible)
        for file in files:
            file_content = await file.read()
            total_pages += count_pdf_pages(file_content)
            
            # Create a temp file that persists after this function finishes
            # (The background task will delete it later)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_content)
                temp_paths.append(tmp.name)

        # 3. QUOTA CHECK (The "Shield")
        # Don't start the expensive AI job if they are out of credits
        if not db_manager.has_enough_quota(user_id, total_pages):
            # Cleanup files since we are aborting
            for p in temp_paths:
                os.remove(p)
            raise HTTPException(status_code=403, detail="Not enough page quota remaining.")

        # 4. Prepare Job Data
        file_names = ", ".join([f.filename for f in files])
        final_case_name = case_name if case_name else files[0].filename

        # Update In-Memory Dict (Fast Status Checks)
        processing_jobs[job_id] = {
            "status": "Starting Batch Extraction...", 
            "filename": file_names,
            "timestamp": datetime.now().isoformat(),
            "total_pages": total_pages,
            "case_name": final_case_name,
            "file_list": [f.filename for f in files],
            "user_id": user_id 
        }

        # 5. Save Initial State to Supabase (Replaces JSON DB)
        # We save it as "Processing" so the user sees it in their dashboard immediately
        db_manager.save_case(
            user_id=user_id,
            job_id=job_id,
            case_name=final_case_name,
            chronology="",
            total_pages=total_pages,
            metadata={"status": "Processing"}
        )

        # 5.1 Save Individual Documents to the documents table
        for file in files:
            # We already have the individual page counts from Fitz
            current_file_pages = count_pdf_pages(await file.read())
            await file.seek(0) # Reset after reading
            
            db_manager.save_document(
                case_id=job_id,
                file_name=file.filename,
                page_count=current_file_pages
            )

        # 6. Trigger Pipeline
        # We pass 'user_id' so the pipeline knows where to save the final result
        background_tasks.add_task(
            run_intelligence_pipeline, 
            job_id, 
            temp_paths, 
            file_names, 
            total_pages, 
            user_id
        )

        # 7. Deduct Quota Immediately
        # We deduct now to prevent "quota racing" (starting 10 jobs at once)
        db_manager.update_user_quota(user_id, total_pages)
        print(f"📉 Quota Deducted: {total_pages} pages for User {user_id}")

        return {"job_id": job_id, "message": f"Pipeline started for {len(files)} files."}

    except Exception as e:
        # Emergency Cleanup if anything above fails
        for p in temp_paths:
            if os.path.exists(p):
                os.remove(p)
        print(f"❌ Endpoint Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/status/{job_id}")
async def check_status(job_id: str, user_id: str): 
    # 1. Fast Path: Check In-Memory (RAM)
    if job_id in processing_jobs:
        return processing_jobs[job_id]

    # 2. Slow Path: Check Supabase (Database)
    try:
        # Fetch the specific case
        cases = db_manager.get_all_cases(user_id)
        job = cases.get(job_id)
        
        if job:
            # --- NEW: Fetch related documents from the documents table ---
            docs_res = db_manager.client.table("documents").select("*").eq("case_id", job_id).execute()
            documents = docs_res.data if docs_res.data else []

            return {
                "status": "Completed" if job.get('chronology_text') else "Processing",
                "case_name": job.get('case_name'),
                "chronology": job.get('chronology_text'),
                "total_pages": job.get('total_pages', 0),
                "metadata": job.get('metadata', {}),
                "documents": documents  # Send the professional list to frontend
            }
    except Exception as e:
        print(f"Status Check Error: {e}")

    return {"status": "Job ID not found"}

@app.get("/quota")
async def get_quota(user_id: str): # <--- Needs user_id to know WHOSE quota to check
    try:
        profile = db_manager.get_user_profile(user_id)
        used = 1500 - profile.get('remaining_quota', 1500)
        return {
            "used": used, 
            "limit": 1500, 
            "remaining": profile.get('remaining_quota', 1500)
        }
    except Exception as e:
        print(f"Quota Error: {e}")
        return {"used": 0, "limit": 1500, "remaining": 1500}

@app.get("/download-report/{job_id}")
async def download_report(job_id: str, user_id: str): # <--- Added user_id for security
    
    # 1. Try In-Memory first (Fastest)
    job = processing_jobs.get(job_id)
    chronology_text = ""
    filename = "Legal_Report"

    # 2. If not in memory, fetch from Supabase
    if not job:
        try:
            cases = db_manager.get_all_cases(user_id)
            job_data = cases.get(job_id)
            if job_data:
                chronology_text = job_data.get('chronology_text', "")
                filename = job_data.get('case_name', "Legal_Report")
        except Exception as e:
            print(f"DB Load Error: {e}")

    # 3. If found in memory, use that text
    if job and not chronology_text:
         chronology_text = job.get("chronology", "")
         filename = job.get("case_name", "Legal_Report")

    # 4. Final Validation
    if not chronology_text:
        raise HTTPException(status_code=404, detail="Report not ready or not found")

    # 5. Generate DOCX
    try:
        # Calls your existing helper function
        doc_stream = create_word_doc_stream(
            chronology_text, 
            job_id, 
            filename
        )
        
        file_bytes = doc_stream.getvalue()
        doc_stream.close()

        # Sanitize filename for HTTP headers
        safe_filename = "".join([c for c in filename if c.isalnum() or c in (' ', '_')]).strip()
        
        return Response(
            content=file_bytes,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                "Content-Disposition": f"attachment; filename={safe_filename}_{job_id[:4]}.docx"
            }
        )
    except Exception as e:
        print(f"Generation Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate report DOCX")

@app.get("/history")
async def get_history(user_id: str):
    """
    Fetches the user's case history from Supabase.
    Returns a dictionary formatted exactly like your old JSON structure.
    """
    try:
        # This calls the function we added to database.py
        # It returns: { "job_id_1": {case_data}, "job_id_2": {case_data} }
        return db_manager.get_all_cases(user_id)
    except Exception as e:
        print(f"❌ History Fetch Error: {e}")
        return {}

@app.post("/chat/{job_id}")
async def chat_with_pdf(job_id: str, payload: dict):
    # 1. Extract inputs
    query = payload.get("query")
    user_id = payload.get("user_id") # <--- Make sure Frontend sends this!
    
    if not query:
        return {"answer": "Please ask a question."}

    # 2. Fetch the "Analysis" (Chronology) from Supabase
    # This gives the AI the high-level summary immediately
    chronology_context = ""
    try:
        # We reuse get_all_cases because it's efficient enough for now
        # Ideally, add a 'get_case(job_id)' method to DatabaseManager later
        cases = db_manager.get_all_cases(user_id)
        case_data = cases.get(job_id, {})
        chronology_context = case_data.get("chronology_text", "")
    except Exception as e:
        print(f"⚠️ Could not fetch chronology: {e}")

    # 3. Fetch the "Evidence" (Safe Text) from S3
    evidence_context = ""
    s3_key = f"safe-text/{job_id}_safe.txt"
    try:
        obj = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
        # Limit to 25k chars to leave room for the Chronology + Response
        evidence_context = obj['Body'].read().decode('utf-8')[:25000]
    except s3_client.exceptions.NoSuchKey:
        pass # It's okay if safe text is missing, we'll use just the analysis

    if not chronology_context and not evidence_context:
        return {"answer": "⚠️ I cannot find any data for this case. It may have been deleted."}

    # 4. STRICT SYSTEM PROMPT
    # We feed both contexts into the prompt clearly separated
    context_block = f"""
    === CASE ANALYSIS ===
    {chronology_context[:15000]} 
    
    === MEDICAL RECORD RAW DATA ===
    {evidence_context}
    """
    
    system_instruction = f"""
    You are a strictly constrained Legal AI Assistant.
    
    YOUR MANDATE:
    1. Answer the user's question using ONLY the provided "MEDICAL RECORD CONTEXT" below.
    2. If the answer is not in the text, state: "I cannot find that information in the provided case documents."
    3. STRICTLY REFUSE all general knowledge questions.
       - If asked "Who is the richest man?", "Write a poem", or "What is the capital of France?", 
       - YOU MUST REPLY: "I am restricted to answering questions strictly related to the case context."
    4. Do not use outside knowledge. If the document doesn't say it, you don't know it.
    
    MEDICAL RECORD CONTEXT:
    {context_block}
    """

    # 5. Ask Groq/Llama
    try:
        chat_response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_instruction},
                {"role": "user", "content": query}
            ],
            temperature=0.1,  # Low temp for accuracy
            max_tokens=600
        )
        return {"answer": chat_response.choices[0].message.content}
        
    except Exception as e:
        print(f"Chat Error: {e}")
        return {"answer": "I encountered an error processing your request."}

@app.delete("/delete/{job_id}")
async def delete_job(job_id: str, user_id: str): # <--- Added user_id param
    # 1. Delete from Supabase (The Database)
    try:
        # Calls the function in database.py
        # RLS policies will ensure users can only delete their own cases
        db_manager.delete_case(job_id, user_id)
    except Exception as e:
        print(f"❌ DB Delete Error: {e}")
        # We continue anyway to clean up S3 and Memory

    # 2. Delete from Memory (The RAM)
    if job_id in processing_jobs:
        del processing_jobs[job_id]

    # 3. Delete from S3 (The Cloud Storage)
    # This prevents "Ghost Files" from costing you money
    try:
        s3_client.delete_object(Bucket=BUCKET_NAME, Key=f"safe-text/{job_id}_safe.txt")
        print(f"🗑️ Deleted S3 Safe Text for {job_id}")
    except Exception as e:
        print(f"⚠️ Could not delete S3 object: {e}")

    return {"message": "Case deleted successfully"}

@app.put("/rename/{job_id}")
async def rename_job(job_id: str, request: RenameRequest):
    new_name = request.new_name
    user_id = request.user_id

    # 1. Update in Supabase
    try:
        # We first need to fetch the existing data to preserve the chronology
        # (Since upsert replaces the row, we don't want to lose the text!)
        current_case = db_manager.get_all_cases(user_id).get(job_id)
        
        if not current_case:
             raise HTTPException(status_code=404, detail="Case not found in database")

        # Save back with the NEW name but OLD chronology/pages
        db_manager.save_case(
            user_id=user_id,
            job_id=job_id,
            case_name=new_name,
            chronology=current_case.get('chronology_text', ''),
            pages=current_case.get('total_pages', 0)
        )
    except Exception as e:
        print(f"❌ Rename DB Error: {e}")
        raise HTTPException(status_code=500, detail="Database rename failed")

    # 2. Update in Memory (RAM)
    # This ensures the user sees the new name instantly without a page refresh
    if job_id in processing_jobs:
        processing_jobs[job_id]["case_name"] = new_name
        
    return {"message": f"Renamed to '{new_name}'"}

@app.post("/append-files/{job_id}")
async def append_files(
    job_id: str, 
    background_tasks: BackgroundTasks, 
    files: List[UploadFile] = File(...),
    user_id: str = Form(...) # <--- NEW: Required for security
):
    # 1. Fetch Current State from Supabase (Source of Truth)
    # We need the old chronology to append to it
    try:
        current_cases = db_manager.get_all_cases(user_id)
        job = current_cases.get(job_id)
        
        if not job:
            raise HTTPException(status_code=404, detail="Case not found in database")
            
        old_chronology = job.get("chronology_text", "")
        # Retrieve existing file list from metadata, defaulting to empty
        current_metadata = job.get("metadata") or {}
        current_file_list = current_metadata.get("file_list", [])
        
    except Exception as e:
        print(f"❌ DB Fetch Error: {e}")
        raise HTTPException(status_code=500, detail="Database error")

    # 2. Process New Files
    new_text_content = ""
    new_filenames = []
    total_new_pages = 0

    for file in files:
        if not file.filename.lower().endswith(".pdf"): continue
        
        # Read content once
        content = await file.read()
        new_filenames.append(file.filename)
        
        # --- PAGE COUNTING ---
        # We use the helper function defined in main.py for consistency
        page_count = count_pdf_pages(content)
        total_new_pages += page_count
        
        # Extract Text (Using your existing process_new_pdf helper)
        # Note: Ensure process_new_pdf is updated to handle bytes/temp files as discussed!
        extracted_text = process_new_pdf(content, file.filename, job_id) 
        new_text_content += f"\n\n--- SUPPLEMENTAL DOCUMENT: {file.filename} ---\n{extracted_text}"

    # 3. QUOTA CHECK (The Shield)
    if total_new_pages > 0:
        if not has_enough_quota(user_id, total_new_pages):
             raise HTTPException(status_code=403, detail="Not enough page quota for these new files.")
        
        # Deduct Quota Immediately
        db_manager.update_user_quota(user_id, total_new_pages)
        print(f"📉 Quota Deducted (Append): {total_new_pages} pages")

    # 4. UPDATE DB STATE
    # We update the file list and status so the UI shows "Updating..."
    updated_metadata = current_metadata.copy()
    updated_metadata["file_list"] = current_file_list + new_filenames
    updated_metadata["status"] = "Updating Chronology..."
    updated_metadata["last_updated"] = datetime.now().isoformat()

    # Save the "In Progress" state to Supabase
    # We keep the old chronology for now until the background task finishes
    db_manager.save_case(
        user_id=user_id,
        job_id=job_id,
        case_name=job.get("case_name"),
        chronology=old_chronology,
        pages=job.get("total_pages", 0) + total_new_pages
    )
    
    # We also need to explicitly update the metadata column since save_case 
    # might not expose it directly in the simplified version I gave you earlier.
    # If your save_case doesn't handle metadata, we do a direct patch:
    try:
        supabase.table("cases").update({
            "metadata": updated_metadata
        }).eq("id", job_id).execute()
    except Exception as e:
        print(f"⚠️ Metadata Update Warning: {e}")

    # 5. Trigger Background Task
    # IMPORTANT: We pass 'user_id' so the background task can save the result!
    background_tasks.add_task(
        run_smart_update, 
        job_id, 
        old_chronology, 
        new_text_content, 
        total_new_pages,
        user_id  # <--- Pass this to the background function
    )

    return {"message": "Supplemental files received. Analysis is updating."}


async def run_smart_update(job_id: str, old_report: str, new_evidence_text: str, new_page_count_int: int, user_id: str):
    try:
        # Update Memory Status (for immediate UI feedback)
        if job_id in processing_jobs:
            processing_jobs[job_id]["status"] = "Merging Supplemental Records..."

        # ---------------------------------------------------------
        # 1. Groq Integration (Your Exact Stricter Prompt)
        # ---------------------------------------------------------
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": """You are a Senior Personal Injury Paralegal.
                
                TASK:
                Merge the EXISTING REPORT with the NEW EVIDENCE into a SINGLE, CONCISE output.
                
                STRICT CONSTRAINTS:
                1. DO NOT repeat the "Existing Report" text.
                2. DO NOT output an "Appendices" section.
                3. Output ONLY the final merged Markdown report.
                4. INTEGRATE the new information into the existing chronology.
                5. MAINTAIN these Markdown sections: # EXECUTIVE CASE SUMMARY, # MEDICAL CHRONOLOGY TABLE, # MEDICAL BILLING LEDGER, # INJURY HIGHLIGHTS, # TOTAL COST.
                6. REWRITE the # EXECUTIVE CASE SUMMARY: It must be a 3-5 sentence overview that incorporates the NEW evidence. (e.g., if surgery occurred in new records, mention it).
                7. CRITICAL: Re-sort the Chronology Table and Billing Ledger into perfect DATE ORDER.
                8. INJURY HIGHLIGHTS: Do not just copy the old ones. Evaluate the new records and add any new critical clinical findings to this list.
                9. Provide the TOTAL_LIST: [x, y, z] data block at the end for the developer.
                
                If the new evidence is minor, the report length should remain similar to the original."""},
                
                {"role": "user", "content": f"EXISTING REPORT:\n{old_report}\n\nNEW EVIDENCE:\n{new_evidence_text}"}
            ],
            temperature=0.1
        )

        updated_report = response.choices[0].message.content
        new_total_val = calculate_total_billing(updated_report)

        # ---------------------------------------------------------
        # 2. S3 Sync (Cloud Native - No Local Files)
        # ---------------------------------------------------------
        s3_key = f"safe-text/{job_id}_safe.txt"
        old_safe_text = ""
        
        # Fetch existing text from S3
        try:
            response = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
            old_safe_text = response['Body'].read().decode('utf-8')
        except Exception:
            print(f"⚠️ S3: No existing safe text found for {job_id}")

        # Combine and Upload
        combined_safe_text = old_safe_text + "\n\n--- SUPPLEMENTAL DATA ---\n" + new_evidence_text
        upload_safe_text_to_s3(job_id, combined_safe_text)

        # ---------------------------------------------------------
        # 3. CRITICAL DATA SAVE STEP (Supabase)
        # ---------------------------------------------------------
        
        # A. Fetch the LATEST state from Supabase
        # We do this to ensure we have the 'file_list' and 'case_name' 
        # that were just updated by the append_files endpoint.
        current_case = db_manager.get_all_cases(user_id).get(job_id)
        
        if not current_case:
            print(f"❌ Critical Error: Case {job_id} disappeared from DB during update.")
            return

        # B. Get Metadata
        current_metadata = current_case.get('metadata') or {}
        
        # Update Status & Timestamp
        current_metadata["status"] = "Completed"
        current_metadata["last_updated"] = datetime.now().isoformat()
        current_metadata["total_billed"] = new_total_val
        
        # C. Handle Page Count
        # NOTE: The 'append_files' endpoint ALREADY added 'new_page_count_int' to the DB.
        # So we just trust the value currently in the database.
        final_page_total = current_case.get('total_pages', 0)
        
        print(f"👉 DEBUG: Final Page Count in DB is {final_page_total}")

        # D. Save Main Fields (Updates the Chronology)
        db_manager.save_case(
            user_id=user_id,
            job_id=job_id,
            case_name=current_case.get('case_name'), # Preserve name
            chronology=updated_report,               # Save new report
            pages=final_page_total                   # Preserve correct count
        )
        
        # E. Patch Metadata (Status & Billing)
        try:
            supabase.table("cases").update({
                "metadata": current_metadata
            }).eq("id", job_id).execute()
        except Exception as e:
            print(f"⚠️ Metadata patch warning: {e}")

        # F. Update Memory (For immediate Dashboard feedback)
        if job_id in processing_jobs:
            processing_jobs[job_id].update({
                "status": "Completed",
                "chronology": updated_report,
                "total_billed": new_total_val,
                "pages": final_page_total,
                "last_updated": datetime.now().isoformat()
            })

        print(f"✅ Case {job_id} merged and saved to Supabase.")

    except Exception as e:
        print(f"❌ Smart Update Failed for {job_id}: {e}")
        if job_id in processing_jobs:
            processing_jobs[job_id]["status"] = f"Update Failed: {str(e)}"